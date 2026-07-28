"""
菌彩野生菌火锅 · 菜品推荐知识库 - 数据处理器

将5个原始数据文件（Excel + Word）解析为结构化文本块，
每个文本块包含自然语言描述文本和元数据，用于后续向量化。
"""

import os
import pandas as pd
from docx import Document
from config import DATA_FILES


def _safe_str(val):
    """安全转换为字符串，NaN返回空字符串"""
    if pd.isna(val):
        return ""
    return str(val).strip()


def _allergen_text(row, cols_labels):
    """生成过敏原信息描述"""
    parts = []
    for col_idx, label in cols_labels:
        val = _safe_str(row[col_idx])
        if val == "有":
            parts.append(f"含{label}")
        elif val == "否":
            parts.append(f"不含{label}")
    if not parts:
        return "无过敏原信息"
    return "，".join(parts)


# ============================================================
# 文件1：菌彩-菜品辣度咸度分级表.xlsx
# ============================================================
def parse_spice_salt_calorie():
    """解析辣度咸度分级表，返回菜品属性块和等级定义块"""
    chunks = []
    path = DATA_FILES["spice_salt_calorie"]
    xls = pd.ExcelFile(path)

    # --- Sheet1: 菜品细化分类表 ---
    df = pd.read_excel(path, sheet_name="菌彩-菜品细化分类表", header=None)
    # Row 1 是表头，Row 2+ 是数据，Col 0-4 是菜品数据
    for idx in range(2, len(df)):
        dish_name = _safe_str(df.iloc[idx, 0])
        if not dish_name:
            continue
        spice = _safe_str(df.iloc[idx, 1])
        salt = _safe_str(df.iloc[idx, 2])
        prop = _safe_str(df.iloc[idx, 3])
        calorie = _safe_str(df.iloc[idx, 4])

        text = (
            f"菜品：{dish_name}。"
            f"辣度：{spice}。"
            f"咸度：{salt}。"
            f"冷热属性：{prop}。"
            f"热量等级：{calorie}。"
        )
        chunks.append({
            "id": f"spice_{idx}",
            "text": text,
            "metadata": {
                "source": "辣度咸度分级表",
                "type": "dish_attribute",
                "dish_name": dish_name,
                "spice_level": spice,
                "salt_level": salt,
                "property": prop,
                "calorie": calorie,
            },
        })

    # --- Sheet2: 辣度等级定义 ---
    df_spice = pd.read_excel(path, sheet_name="菜品辣度等级划分定义", header=None)
    for idx in range(1, len(df_spice)):
        score = _safe_str(df_spice.iloc[idx, 0])
        level = _safe_str(df_spice.iloc[idx, 1])
        definition = _safe_str(df_spice.iloc[idx, 2])
        if not level:
            continue
        text = f"辣度等级定义 - {level}（{score}）：{definition}"
        chunks.append({
            "id": f"spice_def_{idx}",
            "text": text,
            "metadata": {
                "source": "辣度咸度分级表",
                "type": "level_definition",
                "category": "辣度",
                "level": level,
            },
        })

    # --- Sheet3: 盐度等级定义 ---
    df_salt = pd.read_excel(path, sheet_name="菜品盐度等级划分定义", header=None)
    for idx in range(1, len(df_salt)):
        score = _safe_str(df_salt.iloc[idx, 0])
        level = _safe_str(df_salt.iloc[idx, 1])
        definition = _safe_str(df_salt.iloc[idx, 2])
        if not level:
            continue
        text = f"咸度等级定义 - {level}（{score}）：{definition}"
        chunks.append({
            "id": f"salt_def_{idx}",
            "text": text,
            "metadata": {
                "source": "辣度咸度分级表",
                "type": "level_definition",
                "category": "咸度",
                "level": level,
            },
        })

    # --- Sheet4: 热量等级定义 ---
    df_cal = pd.read_excel(path, sheet_name="菜品热量等级划分定义", header=None)
    for idx in range(1, len(df_cal)):
        level = _safe_str(df_cal.iloc[idx, 0])
        score = _safe_str(df_cal.iloc[idx, 1])
        definition = _safe_str(df_cal.iloc[idx, 2])
        if not level:
            continue
        text = f"热量等级定义 - {level}（{score}）：{definition}"
        chunks.append({
            "id": f"cal_def_{idx}",
            "text": text,
            "metadata": {
                "source": "辣度咸度分级表",
                "type": "level_definition",
                "category": "热量",
                "level": level,
            },
        })

    return chunks


# ============================================================
# 文件2：水果过敏原信息.xlsx
# ============================================================
def parse_fruit_allergen():
    """解析水果过敏原信息"""
    chunks = []
    path = DATA_FILES["fruit_allergen"]
    df = pd.read_excel(path, sheet_name="Sheet1")

    for idx, row in df.iterrows():
        risk_level = _safe_str(row.get("风险等级", ""))
        fruit = _safe_str(row.get("水果种类", ""))
        risk_point = _safe_str(row.get("过敏 / 不适风险点", ""))
        advice = _safe_str(row.get("食用建议", ""))
        if not fruit:
            continue
        text = (
            f"水果过敏原信息 - {fruit}："
            f"风险等级：{risk_level}。"
            f"过敏风险：{risk_point}。"
            f"食用建议：{advice}。"
        )
        chunks.append({
            "id": f"fruit_{idx}",
            "text": text,
            "metadata": {
                "source": "水果过敏原信息",
                "type": "fruit_allergen",
                "fruit": fruit,
                "risk_level": risk_level,
            },
        })
    return chunks


# ============================================================
# 文件3：云南菌彩野生菌火锅适合人群.xls
# ============================================================
def parse_suitable_crowd():
    """解析适合人群表"""
    chunks = []
    path = DATA_FILES["suitable_crowd"]
    df = pd.read_excel(path, sheet_name="Sheet1", header=None)

    allergen_cols = [
        (4, "香菜"), (5, "葱"), (6, "蒜"),
        (7, "花生"), (8, "海鲜"), (9, "乳制品"),
    ]

    current_category = ""
    for idx in range(2, len(df)):
        category = _safe_str(df.iloc[idx, 0])
        if category:
            current_category = category
        dish_name = _safe_str(df.iloc[idx, 1])
        if not dish_name:
            continue
        crowd = _safe_str(df.iloc[idx, 2])
        pairing = _safe_str(df.iloc[idx, 3])
        allergen_info = _allergen_text(df.iloc[idx], allergen_cols)

        text = (
            f"菜品：{dish_name}。"
            f"分类：{current_category}。"
            f"适合人群：{crowd}。"
            f"搭配建议：{pairing}。"
            f"过敏原信息：{allergen_info}。"
        )
        chunks.append({
            "id": f"crowd_{idx}",
            "text": text,
            "metadata": {
                "source": "适合人群表",
                "type": "dish_crowd",
                "dish_name": dish_name,
                "category": current_category,
                "suitable_crowd": crowd,
                "pairing": pairing,
                "allergen_info": allergen_info,
            },
        })
    return chunks


# ============================================================
# 文件4：菜品搭配关系.docx
# ============================================================
def parse_dish_pairing():
    """解析菜品搭配关系文档，按方案分块"""
    chunks = []
    path = DATA_FILES["dish_pairing"]
    doc = Document(path)

    current_section = ""
    current_subsection = ""
    current_text_parts = []
    current_section_type = ""

    def _flush():
        nonlocal current_text_parts
        if current_section and current_text_parts:
            full_text = "\n".join(current_text_parts)
            chunks.append({
                "id": f"pairing_{len(chunks)}",
                "text": f"{current_section}\n{full_text}",
                "metadata": {
                    "source": "菜品搭配关系",
                    "type": current_section_type or "combo_plan",
                    "section": current_section,
                },
            })
        current_text_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else "Normal"

        if style == "Heading 1":
            _flush()
            current_section = text
            if "避雷" in text:
                current_section_type = "avoid_combo"
            elif "小贴士" in text or "贴士" in text:
                current_section_type = "tips"
            else:
                current_section_type = "combo_plan"
        elif style == "Heading 2":
            # 子标题作为内容的一部分
            if current_text_parts:
                current_text_parts.append(f"\n【{text}】")
            else:
                current_text_parts.append(f"【{text}】")
        else:
            current_text_parts.append(text)

    _flush()
    return chunks


# ============================================================
# 文件5：13项菌子容易重复、口味冲突.docx
# ============================================================
def parse_mutual_exclusion():
    """解析菌子互斥/口味冲突文档"""
    chunks = []
    path = DATA_FILES["mutual_exclusion"]
    doc = Document(path)

    current_section = ""
    current_text_parts = []

    def _flush():
        nonlocal current_text_parts
        if current_section and current_text_parts:
            full_text = "\n".join(current_text_parts)
            chunks.append({
                "id": f"exclude_{len(chunks)}",
                "text": f"菜品互斥提示 - {current_section}\n{full_text}",
                "metadata": {
                    "source": "菌子重复口味冲突",
                    "type": "mutual_exclusion",
                    "section": current_section,
                },
            })
        current_text_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 检测新的分类标题（以"一、"、"二、"等开头，或包含"门店点单推荐搭配规则"）
        is_title = (
            text.startswith("一、") or text.startswith("二、") or text.startswith("三、")
            or text.startswith("四、") or text.startswith("五、")
            or "门店点单推荐搭配规则" in text
        )

        if is_title:
            _flush()
            current_section = text
        else:
            current_text_parts.append(text)

    _flush()
    return chunks


# ============================================================
# 合并菜品信息：将辣度咸度表和适合人群表的数据合并
# ============================================================
def build_merged_dish_profiles(spice_chunks, crowd_chunks):
    """
    将两个来源的菜品数据按菜名合并，生成完整的菜品档案。
    对于只在其中一个来源出现的菜品，使用已有数据。
    """
    # 构建菜名到数据的映射
    spice_map = {}
    for c in spice_chunks:
        if c["metadata"]["type"] == "dish_attribute":
            name = c["metadata"]["dish_name"]
            spice_map[name] = c["metadata"]

    crowd_map = {}
    for c in crowd_chunks:
        if c["metadata"]["type"] == "dish_crowd":
            name = c["metadata"]["dish_name"]
            crowd_map[name] = c["metadata"]

    # 合并所有菜名
    all_names = set(spice_map.keys()) | set(crowd_map.keys())

    # 菜名别名映射（两个来源中名称不同但指同一菜品）
    aliases = {
        "菌汤锅底": "菌汤生态鸡子母锅",
        "酸汤锅底": "菌汤生态鸡子母锅",
    }

    merged = []
    seen = set()
    for name in sorted(all_names):
        # 处理别名
        resolved_name = aliases.get(name, name)
        if resolved_name in seen:
            continue
        seen.add(resolved_name)

        spice_data = spice_map.get(name, {})
        crowd_data = crowd_map.get(resolved_name, {}) or crowd_map.get(name, {})

        # 合并数据
        dish_name = resolved_name
        category = crowd_data.get("category", spice_data.get("property", ""))
        spice_level = spice_data.get("spice_level", "")
        salt_level = spice_data.get("salt_level", "")
        property_val = spice_data.get("property", "")
        calorie = spice_data.get("calorie", "")
        suitable_crowd = crowd_data.get("suitable_crowd", "")
        pairing = crowd_data.get("pairing", "")
        allergen_info = crowd_data.get("allergen_info", "")

        # 构建自然语言描述
        parts = [f"菜品：{dish_name}。"]
        if category:
            parts.append(f"分类：{category}。")
        if spice_level:
            parts.append(f"辣度：{spice_level}。")
        if salt_level:
            parts.append(f"咸度：{salt_level}。")
        if property_val:
            parts.append(f"冷热属性：{property_val}。")
        if calorie:
            parts.append(f"热量等级：{calorie}。")
        if suitable_crowd:
            parts.append(f"适合人群：{suitable_crowd}。")
        if pairing:
            parts.append(f"搭配建议：{pairing}。")
        if allergen_info:
            parts.append(f"过敏原信息：{allergen_info}。")

        text = "".join(parts)
        merged.append({
            "id": f"dish_profile_{len(merged)}",
            "text": text,
            "metadata": {
                "source": "合并档案",
                "type": "dish_profile",
                "dish_name": dish_name,
                "category": category,
                "spice_level": spice_level,
                "salt_level": salt_level,
                "property": property_val,
                "calorie": calorie,
                "suitable_crowd": suitable_crowd,
                "pairing": pairing,
                "allergen_info": allergen_info,
            },
        })

    return merged


# ============================================================
# 主入口：解析所有文件
# ============================================================
def process_all():
    """解析所有数据文件，返回全部文本块"""
    print("解析辣度咸度分级表...")
    spice_chunks = parse_spice_salt_calorie()
    print(f"  -> {len(spice_chunks)} 条")

    print("解析水果过敏原信息...")
    fruit_chunks = parse_fruit_allergen()
    print(f"  -> {len(fruit_chunks)} 条")

    print("解析适合人群表...")
    crowd_chunks = parse_suitable_crowd()
    print(f"  -> {len(crowd_chunks)} 条")

    print("解析菜品搭配关系...")
    pairing_chunks = parse_dish_pairing()
    print(f"  -> {len(pairing_chunks)} 条")

    print("解析菌子互斥/口味冲突...")
    exclusion_chunks = parse_mutual_exclusion()
    print(f"  -> {len(exclusion_chunks)} 条")

    print("合并菜品完整档案...")
    profile_chunks = build_merged_dish_profiles(spice_chunks, crowd_chunks)
    print(f"  -> {len(profile_chunks)} 条")

    all_chunks = (
        profile_chunks          # 合并后的完整菜品档案
        + spice_chunks          # 原始辣度咸度数据（含等级定义）
        + fruit_chunks          # 水果过敏原
        + crowd_chunks          # 原始适合人群数据
        + pairing_chunks        # 搭配方案
        + exclusion_chunks      # 互斥规则
    )
    print(f"\n总计：{len(all_chunks)} 条文本块")
    return all_chunks


if __name__ == "__main__":
    chunks = process_all()
    print("\n" + "=" * 60)
    print("前5条文本块预览：")
    print("=" * 60)
    for c in chunks[:5]:
        print(f"\n[{c['id']}] type={c['metadata']['type']}")
        print(f"  text: {c['text'][:120]}...")
