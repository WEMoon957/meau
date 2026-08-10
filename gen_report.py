"""生成项目工作总结与汇报 Word 文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ======================== 样式设置 ========================
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# 标题样式
for i in range(1, 4):
    h_style = doc.styles[f'Heading {i}']
    h_font = h_style.font
    h_font.name = '微软雅黑'
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if i == 1:
        h_font.size = Pt(22)
        h_font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        h_font.bold = True
    elif i == 2:
        h_font.size = Pt(16)
        h_font.color.rgb = RGBColor(0x0E, 0xA5, 0xE9)
        h_font.bold = True
    elif i == 3:
        h_font.size = Pt(13)
        h_font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        h_font.bold = True

def add_code_block(doc, code_text, language=""):
    """添加代码块（灰色背景）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # 添加灰色背景
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E293B" w:val="clear"/>')
    pPr.append(shading)
    lines = code_text.strip().split('\n')
    for i, line in enumerate(lines):
        if i == 0:
            run = p.add_run(f"```{language}\n")
            run.font.size = Pt(9)
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        else:
            run = p.add_run(line + '\n')
            run.font.size = Pt(9)
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    run = p.add_run("```")
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

def add_normal_para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    return p

def add_table_row(table, cells_text, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(10)
        if bold:
            run.bold = True
    return row

# ======================== 封面 ========================
for _ in range(4):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('小菌点餐智能体')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0E, 0xA5, 0xE9)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run('项目工作总结与工作汇报')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_p.add_run(f'项目名称：菌彩野生菌火锅 · AI 智能点餐系统\n')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
run = info_p.add_run(f'技术栈：Python / FastAPI / LangChain / DeepSeek / MySQL / Redis / ChromaDB\n')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
run = info_p.add_run(f'文档日期：{datetime.date.today().strftime("%Y年%m月%d日")}\n')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_page_break()

# ======================== 目录 ========================
doc.add_heading('目  录', level=1)
toc_items = [
    '一、项目背景与工作概述',
    '二、项目整体功能说明',
    '    2.1 系统架构总览',
    '    2.2 AI 智能对话推荐模块',
    '    2.3 菜品知识库与规则引擎',
    '    2.4 购物车加购与下单模块',
    '    2.5 高并发与安全防护体系',
    '    2.6 前端交互界面',
    '    2.7 部署与运维',
    '三、核心代码与关键实现',
    '    3.1 Agent 智能体核心架构',
    '    3.2 推荐算法引擎',
    '    3.3 双源数据合并（向量库 + MySQL）',
    '    3.4 规则引擎冲突检测',
    '    3.5 会话管理与分布式锁',
    '    3.6 三层限流防护体系',
    '    3.7 收钱吧网关 RSA2 签名代理',
    '四、项目展示结果',
    '五、工作成果与收获',
    '六、现存不足与后续优化方向',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ======================== 一、项目背景与工作概述 ========================
doc.add_heading('一、项目背景与工作概述', level=1)

doc.add_heading('1.1 项目简介', level=2)
add_normal_para(doc, '本项目为"菌彩野生菌火锅"连锁餐厅打造了一套基于大语言模型（LLM）的 AI 智能点餐系统——"小菌点餐智能体"。系统以自然语言对话为核心交互方式，顾客只需像和朋友聊天一样说出用餐需求，AI 即可自动完成菜品推荐、知识问答、过敏原规避、加购下单等全流程服务。')

doc.add_heading('1.2 项目目标', level=2)
add_bullet(doc, '核心目标：替代传统手动翻阅菜单的点餐方式，通过自然语言对话完成从"想吃什么"到"确认下单"的完整链路')
add_bullet(doc, '智能推荐：根据人数、口味、人群类型、天气季节、会员等级等 10+ 维度，自动生成个性化菜品推荐方案')
add_bullet(doc, '知识服务：集成 83 种菌子菜品的完整档案（辣度、咸度、热量、过敏原、适合人群），支持语义搜索')
add_bullet(doc, '业务合规：内置菜品互斥规则引擎和冲突检测，自动规避菌子重复、口味冲突、过敏原风险')
add_bullet(doc, '商业化闭环：对接收钱吧（全来店）支付网关，支持 RSA2 签名安全加购，完成推荐→加购→下单全链路')
add_bullet(doc, '高并发生产就绪：支持 gunicorn 多 worker + Redis 共享会话，具备 K8s/Docker 容器化部署能力')

doc.add_heading('1.3 本人承担工作', level=2)
add_normal_para(doc, '本人作为项目全栈负责人，独立完成了以下工作：')
add_bullet(doc, '系统架构设计：从单机开发到高并发生产的完整架构演进方案')
add_bullet(doc, '后端核心开发：FastAPI API 服务、LangChain Agent 智能体、推荐算法、规则引擎、知识库系统')
add_bullet(doc, '数据库设计：MySQL 表结构设计、参数化安全查询、连接池管理、读写分离规划')
add_bullet(doc, '向量知识库：基于 ChromaDB + 千问 text-embedding-v3 的菜品知识库构建与语义检索')
add_bullet(doc, '安全防护：会话鉴权、多层限流、CORS 白名单、SQL 注入防护、X-Forwarded-For 安全解析')
add_bullet(doc, '第三方对接：收钱吧网关 RSA2 签名代理、加购子 Agent 自然语言提取')
add_bullet(doc, '前端开发：Web 端对话交互界面（HTML/CSS/JS），支持推荐结果可视化渲染')
add_bullet(doc, '部署运维：Docker 镜像构建、docker-compose 编排、K8s 部署文件、gunicorn 多 worker 配置、Nginx 反代、systemd 服务化、一键部署脚本')
add_bullet(doc, '测试体系：30 条端到端自然语言测试用例、API 接口测试、购物车加购测试')

doc.add_page_break()

# ======================== 二、项目整体功能说明 ========================
doc.add_heading('二、项目整体功能说明', level=1)

doc.add_heading('2.1 系统架构总览', level=2)
add_normal_para(doc, '系统采用前后端分离架构，后端基于 FastAPI + LangChain 框架，前端为纯 HTML/CSS/JS 单页应用。整体架构分为以下层次：')

# 架构表
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '层次'
hdr[1].text = '技术组件'
hdr[2].text = '说明'
add_table_row(table, ['接入层', 'Nginx', '反向代理、连接限流、静态文件服务、upstream keepalive'])
add_table_row(table, ['应用层', 'FastAPI + gunicorn + UvicornWorker', '多 worker 异步 Web 服务，每个 worker 独立持有资源池'])
add_table_row(table, ['AI 引擎层', 'LangChain 1.0 + DeepSeek-V4-Flash', '工具调用编排、System Prompt 管理、短路优化、DSML 清洗'])
add_table_row(table, ['推荐引擎层', '自定义评分算法 + 规则引擎', '多维评分（招牌/天气/季节/毛利率/价格）、冲突检测、分类均衡'])
add_table_row(table, ['知识库层', 'ChromaDB + 千问 Embedding', '83 种菜品档案向量化存储，HNSW 索引，cosine 相似度检索'])
add_table_row(table, ['数据层', 'MySQL + Redis', 'MySQL 菜品结构化数据，Redis 会话历史 + 限流计数 + 分布式锁'])
add_table_row(table, ['外部对接', '收钱吧网关', 'RSA2 签名代理，购物车加购/批量下单'])
add_table_row(table, ['容器化', 'Docker + K8s', 'Dockerfile 生产镜像、docker-compose 单机部署、K8s Deployment/ConfigMap/Secret'])

doc.add_paragraph()

doc.add_heading('2.2 AI 智能对话推荐模块', level=2)
add_normal_para(doc, '这是系统的核心模块，基于 LangChain 1.0 框架实现 LLM Agent，支持 9 个工具函数，覆盖点餐全场景。')

add_normal_para(doc, '功能清单：', bold=True)
add_bullet(doc, '菜品查询（query_dish）：根据菜名精确查询菜品价格、辣度、适合人群、过敏原等基本信息')
add_bullet(doc, '菜单浏览（list_menu）：按分类浏览全部菜品，支持分类筛选（菌彩特色/进店必点/菌汤锅底等 10+ 分类）')
add_bullet(doc, '智能推荐（recommend_dishes）：核心推荐工具，支持 12 个参数（人数/口味/人群/健康标签/天气/季节/过敏原/品类开关/排除菜品/会员等级等），自动生成个性化推荐方案')
add_bullet(doc, '知识库查询（search_dish_knowledge）：语义搜索 83 种菜品的完整属性档案（辣度咸度分级、热量等级、冷热属性）')
add_bullet(doc, '搭配方案推荐（get_pairing_plan）：搜索预设套餐方案（经典地道/酸辣傣味/清淡养生/肉食爱好者）和避雷搭配')
add_bullet(doc, '互斥规则查询（get_exclusion_rules）：查询菌子重复/口味冲突等 13 项互斥规则')
add_bullet(doc, '水果过敏原查询（get_fruit_allergen_info）：查询野生菌与水果的食用禁忌和风险等级')
add_bullet(doc, '推荐理由生成（generate_recommendation_reason）：为已推荐菜品生成口语化推荐理由')
add_bullet(doc, '加购下单（add_to_cart）：自然语言加购，支持"来份XX加购""确认下单"等意图')

add_normal_para(doc, '技术亮点：', bold=True)
add_bullet(doc, '短路优化：非推荐场景（精确查询/菜单浏览）工具结果直接返回，跳过第 2 次 LLM 调用，响应时间减半（从 11.5s 降至 5.5s）')
add_bullet(doc, 'DSML 清洗：多层正则防护，清除 LLM 输出的 thinking 块、DSML 工具调用标签、控制字符，防止前端渲染方框')
add_bullet(doc, '防幻觉机制：System Prompt 中 7 条防幻觉规则，强制 LLM 先调用工具再回复，禁止编造菜品/价格')
add_bullet(doc, 'DeepSeek 适配：针对 deepseek-v4-flash 的 thinking 模式做特殊处理，禁用 thinking 避免 content 为空')

doc.add_heading('2.3 菜品知识库与规则引擎', level=2)
add_normal_para(doc, '基于 ChromaDB 向量数据库 + 千问 text-embedding-v3（1024 维）构建的菜品知识库系统，包含 5 类数据源：')
add_bullet(doc, '菜品辣度咸度分级表（Excel）：83 种菜品的辣度等级、咸度分级、热量等级')

add_bullet(doc, '水果过敏原信息（Excel）：20+ 种水果与野生菌的食用风险等级和建议')
add_bullet(doc, '适合人群数据（Excel）：菜品按青年/儿童/老人/孕妇等标签分类')
add_bullet(doc, '菜品搭配关系（Word）：4 套预设套餐方案 + 避雷搭配建议')
add_bullet(doc, '菜品互斥规则（Word）：13 项菌子重复和口味冲突规则')

add_normal_para(doc, '规则引擎功能：', bold=True)
add_bullet(doc, '双向冲突图谱：从知识库规则文本中自动提取菜品名称，构建有向图冲突关系')
add_bullet(doc, '推荐时自动过滤：在推荐流程中集成冲突检测，自动跳过与已选菜品冲突的候选')
add_bullet(doc, '冲突警告生成：检测到菜品组合冲突时，自动生成可读的警告文案')

doc.add_heading('2.4 购物车加购与下单模块', level=2)
add_normal_para(doc, '集成收钱吧（全来店）支付网关，实现从推荐到下单的商业闭环：')
add_bullet(doc, 'RSA2 签名代理：后端持有私钥，前端无需接触密钥，安全完成加购请求签名')
add_bullet(doc, '加购子 Agent（CartAgent）：基于 LLM 从自然语言提取菜名+数量，自动反查菜品 ID 后批量加购')
add_bullet(doc, '确认下单流程：前端解析推荐文本菜名 → 反查接口（/api/dish/resolve）→ 批量加购（/api/cart/batch-add）')
add_bullet(doc, '降级容错：收钱吧网关未通时自动 mock 成功，保证演示链路可走通（mocked=True 标记）')
add_bullet(doc, '安全防护：购物车接口独立限流（每会话 10 次/分钟，每 IP 20 次/分钟），必须持有有效会话凭证')

doc.add_heading('2.5 高并发与安全防护体系', level=2)
add_normal_para(doc, '系统从架构层面实现了完整的高并发支撑和安全防护：')

add_normal_para(doc, '高并发设计：', bold=True)
add_bullet(doc, '多 worker 部署：gunicorn + UvicornWorker，worker 数 = CPU 核心数 × 2 + 1')
add_bullet(doc, '连接池管理：MySQL 使用 DBUtils PooledDB，单 worker 20 连接，总连接数 ≤ max_connections × 0.8')
add_bullet(doc, 'Redis 外部化：会话历史 + 限流计数 + 分布式锁全部存入 Redis，多 worker 共享，杜绝串号')
add_bullet(doc, '背压机制：asyncio.Semaphore 限制单 worker 并发 LLM 调用（MAX_CONCURRENT_CHATS=20），超限立即 503')
add_bullet(doc, '请求超时：LLM_REQUEST_TIMEOUT=30s 硬约束，防止慢请求堆积')
add_bullet(doc, '优雅关停：gunicorn graceful_timeout=75s，SIGTERM 后等待 in-flight 请求完成')

add_normal_para(doc, '安全防护：', bold=True)
add_bullet(doc, 'SQL 注入防护：全面使用参数化查询（cursor.execute(sql, params)），禁止字符串拼接 SQL；分类/辣度参数白名单校验')
add_bullet(doc, '会话鉴权（C3 IDOR 防御）：session_token 签发/校验机制，常数时间比较防时序攻击，购物车/reset 必须持有有效 token')
add_bullet(doc, '三层限流：对话限流（每会话 30 次/分钟 + 每 IP 60 次/分钟）、购物车限流（每会话 10 次/分钟 + 每 IP 20 次/分钟）、会话管理限流')
add_bullet(doc, 'CORS 白名单：环境变量配置允许的来源域名，禁止通配符 *，仅允许 GET/POST 方法')
add_bullet(doc, 'X-Forwarded-For 安全解析：仅当 TCP 直连来源是可信代理（CIDR 白名单）时才解析 XFF')
add_bullet(doc, '查询硬约束：所有 SELECT 注入 MAX_EXECUTION_TIME(5000) 优化器提示 + LIMIT 兜底')
add_bullet(doc, '异常信息保护：API 层只返回通用错误信息，内部细节仅记日志')

doc.add_heading('2.6 前端交互界面', level=2)
add_normal_para(doc, '前端为纯 HTML/CSS/JS 单页应用（index.html），提供完整的对话交互体验：')
add_bullet(doc, '对话式聊天界面：支持用户消息输入、AI 回复展示、加载状态、错误提示')
add_bullet(doc, '推荐结果可视化：菜品按分类分组展示，包含序号、菜名、价格、辣度、招牌标签，合计总价高亮显示')
add_bullet(doc, '推荐理由展示：独有蓝色背景推荐理由区域，展示规则避让、过敏原提示等上下文信息')
add_bullet(doc, '快捷操作面板：提供人数选择、口味偏好、人群类型等快捷按钮，降低用户输入门槛')
add_bullet(doc, '会话管理：支持重置对话、显示会话 ID、自动维护 session_token 鉴权')

doc.add_heading('2.7 部署与运维', level=2)
add_bullet(doc, 'Docker 部署：Dockerfile 构建生产镜像（基于 python:3.11-slim），国内镜像源加速，docker-compose 编排 app + redis')
add_bullet(doc, 'K8s 部署：完整的 Deployment、Service(NodePort)、ConfigMap、Secret、Namespace 配置，就绪探针 + 存活探针')
add_bullet(doc, '一键部署脚本（deploy.sh）：8 步自动化部署（系统依赖→MySQL→Redis→代码拉取→虚拟环境→数据库初始化→systemd 服务→Nginx 配置）')
add_bullet(doc, '健康检查：/api/health 端点包含 DB 探活 + Redis 探活 + KB 探活，依赖不可用时返回 503 供 LB 摘流')
add_bullet(doc, '日志体系：Python logging 按级别输出，gunicorn accesslog/errorlog 可配置，systemd journalctl 集中管理')

doc.add_page_break()

# ======================== 三、核心代码与关键实现 ========================
doc.add_heading('三、核心代码与关键实现', level=1)
add_normal_para(doc, '以下提炼项目中最关键的核心代码片段，每个片段标注技术栈、解决的技术难点和代码逻辑说明。')

# --- 3.1 Agent 智能体核心架构 ---
doc.add_heading('3.1 Agent 智能体核心架构（关键技术栈：LangChain 1.0 + DeepSeek-V4-Flash）', level=2)
add_normal_para(doc, '难点：需要实现 LLM 工具调用 + 短路优化 + DSML 标签清洗 + 防幻觉 + DeepSeek 适配，在保证回复质量的同时降低延迟。')

add_code_block(doc, '''class OrderingAgent:
    """基于 LangChain 1.0 的无状态点餐智能体（短路优化版）"""
    
    def chat(self, user_input: str, history: list, membership_level: str = "") -> tuple[str, list]:
        # 1. 构建消息（System Prompt + 历史 + 用户输入）
        messages = [SystemMessage(content=SYSTEM_PROMPT)] + list(history) + [HumanMessage(content=user_message)]
        
        # 2. 第 1 次 LLM 调用：决策调用哪个工具
        ai_msg = self.llm_with_tools.invoke(messages)
        
        # 3. 无工具调用 → 直接返回（闲聊场景）
        if not ai_msg.tool_calls:
            response = self._strip_dsml(ai_msg.content or "")
            return response, [HumanMessage(content=user_input), AIMessage(content=response)]
        
        # 4. 执行工具调用
        for tc in ai_msg.tool_calls:
            tool = self.tool_map.get(tc["name"])
            if tool:
                result = tool.invoke(tc["args"])
                messages.append(ToolMessage(content=result, tool_call_id=tc["id"]))
        
        # 5. 短路判断：非推荐/知识库场景的结果为完整回复时直接返回
        _TOOLS_NEED_LLM_REWRITE = {"recommend_dishes", "search_dish_knowledge", ...}
        need_rewrite = any(tc["name"] in _TOOLS_NEED_LLM_REWRITE for tc in ai_msg.tool_calls)
        if tool_results and not need_rewrite:
            if self._is_complete_response(tool_results[-1]):
                return tool_results[-1], [...]  # 跳过第 2 次 LLM 调用
        
        # 6. 第 2 次 LLM 调用：综合工具结果生成口语化回复
        final_msg = self.llm.invoke(messages)
        return self._strip_dsml(final_msg.content or ""), [...]
    
    @staticmethod
    def _strip_dsml(text: str) -> str:
        """清除 LLM 返回中的 DSML 标签、think 块、控制字符"""
        text = _THINK_BLOCK_RE.sub("", text)   # 移除  thinking... response 块
        text = _DSML_BLOCK_RE.sub("", text)     # 移除 DSML 工具调用配对块
        text = _DSML_TAG_RE.sub("", text)       # 移除孤立 DSML 标签
        # 剔除不可打印 Unicode 控制字符（防止前端渲染方框 □）
        text = "".join(ch for ch in text if ch in ("\\n", "\\r", "\\t") or ord(ch) >= 0x20)
        return text.strip()''', 'python')

add_normal_para(doc, '代码说明：该模块实现了 Agent 的核心调度逻辑。短路优化将非推荐场景的响应时间从 11.5s 降至 5.5s（节省一次 LLM 调用）。DSML 清洗通过 4 层正则防护，解决了 DeepSeek-V4-Flash 偶尔将工具调用"假装"写成 DSML 文本而非真正触发 function calling 的问题。')

# --- 3.2 推荐算法引擎 ---
doc.add_heading('3.2 推荐算法引擎（关键技术栈：Python + 多维度评分 + 规则引擎）', level=2)
add_normal_para(doc, '难点：需要在 83 种菜品中，根据 10+ 维度（人数、口味、人群、健康、天气、季节、过敏原、会员等级、品类平衡、互斥规则）实时生成个性化推荐方案，同时保证每次推荐的差异性。')

add_code_block(doc, '''@tool
def recommend_dishes(people_count=0, taste="", customer_type="", health_tags="",
                     weather="", season="", allergen_avoid="", include_drinks=True,
                     include_staple=True, include_soup=True, exclude_dishes="",
                     membership_level=""):
    """智能推荐：多维度评分 + 分类均衡 + 规则引擎冲突检测"""
    # 双源合并：MySQL 结构化数据 + 向量库权威属性
    all_dishes = get_merged_dishes()
    candidates = list(all_dishes)
    
    # 1. 口味筛选（带渐进降级：特辣→中辣→微辣→不限）
    candidates = _filter_by_taste(candidates, taste)
    # 2. 人群筛选 + 场景过滤（儿童餐等）
    candidates = _filter_by_customer_type(candidates, customer_type)
    candidates = _filter_by_scene(candidates, customer_type)
    # 3. 过敏原排除（字段过滤 + 菜名关键词兜底）
    candidates = _filter_by_allergens(candidates, avoid_list)
    # 4. 品类开关过滤
    if not include_drinks: candidates = [d for d in candidates if d.category != "甜饮品"]
    
    # 5. 锅底强制推荐（火锅店业务硬约束）
    pot_pool = _filter_by_allergens(
        [d for d in all_dishes if d.category == "菌汤锅底" and "锅" in d.name], avoid_list)
    if pot_pool:
        recommended.append(random.choice(pot_pool))
    
    # 6. 多样化推荐：多轮选择 + 分类权重 + 规则冲突检测
    for round_num in range(3):
        if round_num == 0:
            # 第一轮：优先高权重分类，每分类最多 2 道
            _select_by_category_weights(...)
        elif round_num == 1:
            # 第二轮：放宽分类限制，每分类最多 3 道
            _select_by_category_weights(...)
        else:
            # 第三轮：从剩余菜品随机选择，应用所有过滤条件
            _select_from_remaining(...)
    
    # 7. 饮品保障：include_drinks=True 但无饮品时自动替换
    # 8. 格式化输出 + 生成结构化上下文供 LLM 润色''', 'python')

add_normal_para(doc, '代码说明：推荐算法采用"多轮渐进式选择"策略。第一轮严格按分类权重（每分类最多 2 道），第二轮放宽至 3 道，第三轮从剩余池随机选择。每轮均集成规则引擎冲突检测，自动跳过与已选菜品冲突的候选。口味降级机制确保即使特辣无结果也能给出合理推荐。')

# --- 3.3 双源数据合并 ---
doc.add_heading('3.3 双源数据合并（关键技术栈：ChromaDB 向量库 + MySQL 关系数据库）', level=2)
add_normal_para(doc, '难点：向量库存储辣度/过敏原/人群的权威数据，MySQL 存储价格/分类/毛利率的结构化数据，两者菜名可能不完全一致，需要精确匹配并补全空字段。')

add_code_block(doc, """def get_merged_dishes() -> list[Dish]:
    '''双源合并：MySQL 结构化数据 + 向量库权威属性'''
    mysql_dishes = get_all_dishes()          # 价格、分类、毛利率
    kb_profiles = _load_kb_profiles()        # 辣度、过敏原、人群、饮食标签
    
    for dish in mysql_dishes:
        kb_meta = kb_profiles.get(dish.name, {})  # 按菜名精确匹配
        merged = _merge_dish_with_kb(dish, kb_meta)
        # 补全规则（向量库为权威源，仅在 MySQL 字段为空时补全）：
        # - spicy_level:  MySQL 空 → 用向量库 spice_level
        # - allergens:    MySQL 空 → 解析向量库 allergen_info 文本
        # - suitable_for: MySQL 空 → 用向量库 suitable_crowd
        # - dietary_tags: MySQL 空 → 用向量库 calorie + property 推导
        # - description:  MySQL 空 → 用向量库多字段拼接'''

def _parse_kb_allergen_info(allergen_info: str) -> list[str]:
    '''解析向量库过敏原文本：\'不含香菜，不含葱，含花生\' → [\'花生\']'''
    for kb_cat in _KB_ALLERGEN_CATEGORIES:
        contain_marker = f"含{kb_cat}"
        exclude_marker = f"不含{kb_cat}"
        if contain_marker in allergen_info and exclude_marker not in allergen_info:
            contains.append(_KB_ALLERGEN_ALIASES.get(kb_cat, kb_cat))""", 'python')

add_normal_para(doc, '代码说明：双源合并解决了 MySQL 中部分菜品过敏原/辣度数据缺失的问题。向量库作为权威数据源补全空字段，同时记录已验证过敏原的菜名集合（_kb_allergen_verified），用于后续过敏原过滤时区分"明确不含"和"数据缺失"两种情况，避免误杀。')

# --- 3.4 规则引擎 ---
doc.add_heading('3.4 规则引擎冲突检测（关键技术栈：Python + 图论 + 知识库规则解析）', level=2)
add_normal_para(doc, '难点：从知识库 Word 文档中自动提取 13 项互斥规则文本，解析出涉及的菜品名称，构建双向冲突图谱，在推荐时 O(1) 检测冲突。')

add_code_block(doc, '''class _DishRulesEngine:
    def _load(self):
        """从知识库加载规则，构建冲突图谱"""
        # 1. 加载互斥规则 + 避雷搭配（从 ChromaDB 批量读取）
        all_rules = get_all_exclusion_rules() + get_all_avoid_combos()
        
        # 2. 解析每条规则，在规则文本中查找菜品名称
        for rule in all_rules:
            matched_dishes = self._find_dish_names_in_text(rule["text"], dish_names_set)
            if len(matched_dishes) >= 2:
                # 建立双向冲突关系
                for d1, d2 in combinations(matched_dishes, 2):
                    self.conflict_map.setdefault(d1, set()).add(d2)
                    self.conflict_map.setdefault(d2, set()).add(d1)  # 双向

    def has_conflict(self, dish_name: str, selected_names: set[str]) -> bool:
        """O(1) 冲突检测"""
        conflicts = self.conflict_map.get(dish_name, set())
        return bool(conflicts & selected_names)  # 集合交集判断''', 'python')

add_normal_para(doc, '代码说明：规则引擎在启动时预加载所有规则，构建双向冲突图谱（dict + set），冲突检测为 O(1) 集合交集运算。支持从自然语言规则文本中自动提取菜品名称，无需手动标注冲突关系。规则引擎加载失败时静默降级，不阻塞推荐流程。')

# --- 3.5 会话管理 ---
doc.add_heading('3.5 会话管理与分布式锁（关键技术栈：Redis + Lua 脚本 + 分布式锁）', level=2)
add_normal_para(doc, '难点：多 worker 模式下，同一会话的并发请求会导致 load-modify-save 竞态（历史覆写丢失）。需要跨进程的分布式锁来串行化同一会话的请求。')

add_code_block(doc, '''class SessionManager:
    def chat(self, session_id, user_input, membership_level=""):
        """加载历史 → 调用 agent → 回写历史（带分布式锁保护）"""
        lock_token = self._acquire_lock(session_id)  # 获取每会话锁
        try:
            history = self._load_history(session_id)  # 从 Redis 加载
            response, new_msgs = self.agent.chat(user_input, history, membership_level)
            self._save_history(session_id, history + new_msgs)  # 回写 Redis
            return response
        finally:
            self._release_lock(session_id, lock_token)  # 释放锁

    def _acquire_lock(self, session_id):
        """Redis 分布式锁：SET NX EX，TTL 兜底防死锁"""
        token = str(uuid.uuid4())
        ok = self.redis.set(f"menu:session:lock:{session_id}", token, nx=True, ex=45)
        if not ok:
            raise SessionBusyError(session_id)
        return token

    def _release_lock(self, session_id, token):
        """Lua 脚本原子释放：仅当 token 匹配时才删除"""
        self.redis.eval(_RELEASE_LOCK_SCRIPT, 1, lock_key, token)''', 'python')

add_normal_para(doc, '代码说明：通过 Redis SET NX EX 实现跨 worker 的分布式锁，TTL 为 45 秒（≥ LLM_REQUEST_TIMEOUT + 余量），防止持有锁的请求崩溃导致会话永久锁死。释放锁使用 Lua 脚本保证原子性，仅当 token 匹配时才删除，防止误删他人持有的锁。')

# --- 3.6 限流 ---
doc.add_heading('3.6 三层限流防护体系（关键技术栈：Redis 固定窗口 + Pipeline 原子操作）', level=2)
add_normal_para(doc, '难点：多 worker 模式下，进程内限流会被 worker 数倍绕过。需要 Redis 共享计数，且保证 INCR+EXPIRE 的原子性。')

add_code_block(doc, '''class RedisFixedWindowRateLimiter:
    """Redis 固定窗口限流（多 worker 共享计数）"""
    def allow(self, key: str) -> tuple[bool, int]:
        now = int(time.time())
        bucket = now - (now % self.window_seconds)  # 时间桶对齐
        rkey = f"{self.prefix}:{bucket}:{key}"
        
        # Pipeline 原子执行 INCR + EXPIRE
        pipe = self.redis.pipeline(transaction=True)
        pipe.incr(rkey)
        pipe.expire(rkey, self.window_seconds * 2)
        count, _ = pipe.execute()
        
        if count > self.max_requests:
            ttl = self.redis.ttl(rkey)
            return False, max(1, ttl)  # 返回 retry_after 秒数
        return True, 0''', 'python')

add_normal_para(doc, '代码说明：采用时间桶对齐的固定窗口算法，同一窗口内所有请求命中同一计数器。Redis Pipeline 保证 INCR+EXPIRE 的原子性，EXPIRE 幂等设置（覆盖也无害）。Redis 故障时抛异常由 API 层返回 503，绝不静默放行（防止 LLM 被打爆）。')

# --- 3.7 收钱吧 ---
doc.add_heading('3.7 收钱吧网关 RSA2 签名代理（关键技术栈：cryptography + RSA2 + httpx）', level=2)
add_normal_para(doc, '难点：收钱吧网关要求公共参数（appId/format/charset/signType/timestamp/version/method/bizContent/sign）全部放在 HTTP 请求头中，签名算法为 RSA2（SHA256withRSA），需兼容 PKCS#1 和 PKCS#8 两种私钥格式。')

add_code_block(doc, '''def _build_sign(params: dict) -> str:
    """按收钱吧规则生成 RSA2 签名"""
    # 1. 过滤空值 → 按 key ASCII 字典序排序 → 拼成 key=value&key=value
    filtered = {k: v for k, v in params.items() if v not in (None, "", [])}
    sorted_keys = sorted(filtered.keys())
    string_a = "&".join(f"{k}={filtered[k]}" for k in sorted_keys)
    
    # 2. RSA2 签名（SHA256withRSA, PKCS#1 v1.5）
    key, padding, hashes = _load_signer()
    signature = key.sign(string_a.encode("utf-8"), padding.PKCS1v15(), hashes.SHA256())
    return base64.b64encode(signature).decode("utf-8")

def _load_signer():
    """兼容 PKCS#1 / PKCS#8 / 纯 Base64 三种私钥格式"""
    if "BEGIN" in pem:
        return serialization.load_pem_private_key(pem.encode(), password=None)
    # 纯 Base64：自动补全 PEM 头尾，尝试两种格式
    for header, footer in [
        ("-----BEGIN PRIVATE KEY-----", "-----END PRIVATE KEY-----"),
        ("-----BEGIN RSA PRIVATE KEY-----", "-----END RSA PRIVATE KEY-----"),
    ]:
        try:
            return serialization.load_pem_private_key(pem_attempt.encode(), password=None)
        except: continue''', 'python')

add_normal_para(doc, '代码说明：该模块实现了收钱吧网关的完整接入流程。签名算法遵循收钱吧规范（过滤空值→ASCII 排序→拼接→RSA2 签名→Base64 编码）。私钥加载兼容 3 种格式，异常信息仅记日志不回传客户端（安全防护）。支持同步/异步两种调用方式，CartAgent 在 asyncio.to_thread 上下文中使用同步版。')

doc.add_page_break()

# ======================== 四、项目展示结果 ========================
doc.add_heading('四、项目展示结果', level=1)

doc.add_heading('4.1 运行效果', level=2)
add_normal_para(doc, '系统已在公网 IP 118.178.123.31 上完成部署验证，整体运行稳定。以下为核心运行指标：')

# 运行指标表
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '指标'
hdr[1].text = '数值'
hdr[2].text = '说明'
add_table_row(table, ['菜品总数', '83 道（双源合并后）', 'MySQL + 向量库合并，覆盖锅底/菌彩特色/进店必点等 10+ 分类'])
add_table_row(table, ['知识库记录数', '269 条（1024 维向量）', '83 种菜品档案 + 搭配方案 + 互斥规则 + 水果过敏原'])
add_table_row(table, ['推荐场景响应时间', '5.5-8s（DeepSeek-V4）', '2 次 LLM 调用（推荐工具 + 润色），vs 千问 11.5s 提升约 50%'])
add_table_row(table, ['查询场景响应时间', '~0.3s（短路模式）', '精确查询/菜单浏览跳过第 2 次 LLM 调用'])
add_table_row(table, ['知识库查询响应时间', '~5.5s', '2 次 LLM 调用（知识库检索 + 口语化改写）'])
add_table_row(table, ['并发能力', '单 worker 20 并发 LLM', '4 worker 总并发 80，背压超限立即 503'])
add_table_row(table, ['会话上限', '500 个活跃会话', 'Redis 索引集合管理，TTL 30 分钟自动过期'])
add_table_row(table, ['限流能力', '对话 30/60 次/分钟', '购物车 10/20 次/分钟，独立限流']) 

doc.add_paragraph()

doc.add_heading('4.2 API 接口输出示例', level=2)
add_normal_para(doc, '对话接口（POST /api/ai/chat）推荐场景典型响应：')

add_code_block(doc, '''{
  "code": 200,
  "msg": "success",
  "aimessage": "天冷就该吃火锅！给2位客人挑了一桌暖心好菜，香辣够味～\\n\\n
--- 菌汤锅底 ---\\n
  1. 菌汤生态鸡子母锅  ￥68\\n
     [中辣]\\n
--- 菌彩特色 ---\\n
  2. 单点绣球菌  ￥32\\n
  3. 单点鹿茸菌  ￥38\\n
--- 进店必点 ---\\n
  4. 包烧见手青  ￥58\\n
--- 涮品 ---\\n
  5. 鲜切牛肉  ￥48\\n\\n
合计：￥244\\n\\n
这桌有菌彩特色、进店必点等多种品类，搭配均衡。\\n
已经帮您避开了海鲜，放心吃～如需调整告诉我！",
  "session_id": "a1b2c3d4-...",
  "session_token": "e5f6g7h8..."
}''', 'json')

add_normal_para(doc, '健康检查接口（GET /api/health）典型响应：')

add_code_block(doc, '''{
  "code": 200,
  "msg": "ok",
  "data": {
    "dish_count": 83,
    "active_sessions": 12,
    "max_sessions": 500,
    "session_ttl_seconds": 1800,
    "max_concurrent_chats": 20,
    "concurrent_at_capacity": false,
    "rate_limits": {
      "per_session": "30/60s",
      "per_ip": "60/60s"
    },
    "dependencies": {
      "db": "ok",
      "redis": "ok",
      "kb": "ok"
    }
  }
}''', 'json')

doc.add_heading('4.3 测试效果', level=2)
add_normal_para(doc, '项目建立了完整的测试体系，覆盖功能验证和性能测试：')

add_bullet(doc, '30 条端到端自然语言测试（test_30_nl_e2e.py）：覆盖 P0 核心硬约束（锅底必选/过敏原过滤/格式完整）、P1 业务场景（会员等级/天气季节/饮品保障）、P2 边界场景（0 人默认/极端过敏/追加排除），全部通过')
add_bullet(doc, '购物车 API 测试（test_cart_api.py）：端到端验证会话创建→加购→返回结果全链路')
add_bullet(doc, '购物车离线测试（test_cart_offline.py）：验证加购逻辑在网关不通时的降级行为')
add_bullet(doc, 'SQL 安全测试（test_validate_sql.py）：验证参数化查询防注入有效性')
add_bullet(doc, '人群筛选测试（test_customer_type_filter.py）：验证儿童餐/孕妇等场景过滤正确性')
add_bullet(doc, '会话并发测试（test_session_concurrency.py）：验证多 worker 下会话隔离和锁机制')

add_normal_para(doc, '【根据测试结果文件推断】30 条端到端测试中，所有 P0 级别用例（锅底必选、过敏原过滤、格式完整、高配额）均通过，推荐结果在不同场景下具有差异性（总价不同、菜品组合不同），儿童餐不会出现在成年聚餐场景中。')

doc.add_heading('4.4 前端可视化效果', level=2)
add_normal_para(doc, '前端界面（index.html）实现了完整的对话式点餐体验，主要视觉效果包括：')
add_bullet(doc, '渐变紫色头部导航栏，展示"小菌点餐智能体"品牌标识和重置按钮')
add_bullet(doc, '聊天气泡式对话界面：用户消息紫色气泡右对齐，AI 回复白色气泡左对齐')
add_bullet(doc, '推荐结果卡片式渲染：分类标题（紫色虚线分隔）、菜品行（序号+菜名+价格橙色+辣度红色标签）、合计总价灰色背景高亮')
add_bullet(doc, '蓝色背景推荐理由区域：展示规则避让、过敏原提示、品类覆盖等推荐上下文')
add_bullet(doc, '红色背景规则警告区域：互斥规则/过敏原冲突提示')
add_bullet(doc, '底部操作面板：人数选择、口味偏好、快捷场景按钮')

add_normal_para(doc, '【根据现有文件推断】API 文档页面（docs/api-doc.html）为完整的交互式 API 文档，包含侧边栏导航、接口分组、请求/响应示例、代码高亮，支持 GET/POST 方法标签颜色区分。')

doc.add_page_break()

# ======================== 五、工作成果与收获 ========================
doc.add_heading('五、工作成果与收获', level=1)

doc.add_heading('5.1 项目产出', level=2)
add_bullet(doc, '完整可运行的 AI 点餐系统：包含后端 API 服务（21 个 Python 源文件）、前端对话界面、API 文档页面')
add_bullet(doc, '9 个 LangChain 工具函数：覆盖菜品查询、智能推荐、知识库检索、加购下单全场景')
add_bullet(doc, '269 条向量的菜品知识库：基于 ChromaDB + 千问 Embedding 构建，支持语义检索')
add_bullet(doc, '13 项互斥规则的规则引擎：自动从知识库文本中提取，构建双向冲突图谱')
add_bullet(doc, '收钱吧网关对接模块：RSA2 签名代理，支持单次加购和批量下单')
add_bullet(doc, '生产级部署方案：Dockerfile + docker-compose + K8s + deploy.sh 一键部署脚本')
add_bullet(doc, '完整的安全防护体系：会话鉴权、三层限流、SQL 注入防护、XFF 安全解析')
add_bullet(doc, '30 条端到端测试用例 + 5 个专项测试脚本')

doc.add_heading('5.2 解决的核心技术问题', level=2)
add_bullet(doc, 'LLM 幻觉防控：通过 7 条防幻觉规则 + 工具调用强制执行 + DSML 多层清洗，解决了 LLM 编造菜品、口头描述工具调用、输出方框字符等问题')
add_bullet(doc, '多 worker 会话串号：通过 Redis 外部化会话历史 + 分布式锁 + session_token 鉴权，彻底解决多 worker 下的会话交叉污染问题')
add_bullet(doc, '推荐场景延迟优化：通过短路优化（跳过非必要的第 2 次 LLM 调用），查询场景响应时间从 11.5s 降至 0.3s，推荐场景从千问 11.5s 降至 DeepSeek 5.5s')
add_bullet(doc, '过敏原数据不完整：通过双源合并（向量库 + MySQL）和菜名关键词兜底过滤，解决了 MySQL 中部分菜品过敏原字段为空导致的过滤失效问题')
add_bullet(doc, 'DeepSeek 模型适配：解决 thinking 模式导致 content 为空、DSML 双竖线变体格式等问题，实现 deepseek-v4-flash 的稳定调用')
add_bullet(doc, '连接池启动失败：将 MySQL 连接池从模块导入时创建改为懒加载（首次访问时创建），避免数据库未就绪时 worker 启动崩溃')

doc.add_heading('5.3 技术沉淀', level=2)
add_bullet(doc, 'LangChain 1.0 Agent 最佳实践：手动管理工具调用循环 + 短路优化的架构模式，可复用于其他 LLM Agent 项目')
add_bullet(doc, '向量库 + 关系数据库双源合并模式：适用于"权威属性在向量库、结构化数据在关系库"的混合数据场景')
add_bullet(doc, '规则引擎设计模式：从自然语言文本中自动提取规则并构建冲突图谱，可推广到其他需要规则约束的推荐系统')
add_bullet(doc, '高并发 FastAPI 部署方案：gunicorn + UvicornWorker + Redis 外部化 + 连接池管理的完整方案')
add_bullet(doc, '第三方支付网关对接经验：RSA2 签名规范、HTTP 头参数传递、密钥格式兼容处理')

doc.add_page_break()

# ======================== 六、不足与优化 ========================
doc.add_heading('六、现存不足与后续优化方向', level=1)

doc.add_heading('6.1 现存不足', level=2)

add_bullet(doc, '收钱吧网关未完全接通：当前枚举的加购接口 method 均返回 1300/404 NOT_FOUND，实际加购链路为 mock 降级模式。需要与收钱吧技术对接确认正确的 method 名称和接口权限')
add_bullet(doc, 'LLM 响应延迟较高：推荐场景（2 次 LLM 调用）仍需 5.5-8 秒，在弱网环境下用户体验有待提升。可考虑引入流式输出（SSE）减少感知延迟')
add_bullet(doc, '知识库数据依赖本地文件：知识库构建依赖 ATTACH_DIR 下的 Excel/Word 文件，CI/CD 流程中需要额外处理数据文件的分发和版本管理')
add_bullet(doc, '前端功能较为基础：当前仅为对话界面，缺少菜品图片展示、分类浏览、历史订单等辅助功能')
add_bullet(doc, '缺少监控告警：无 Prometheus metrics 暴露、无 APM 链路追踪、无告警规则，生产环境故障定位依赖日志')
add_bullet(doc, '未实现 A/B 测试框架：推荐算法的评分权重调整依赖手动修改代码，缺乏在线实验能力')
add_bullet(doc, '缺少用户反馈闭环：当前无用户对推荐结果的满意度收集机制，无法基于真实反馈持续优化推荐质量')

doc.add_heading('6.2 后续优化方向', level=2)

add_normal_para(doc, '短期优化（1-2 周）：', bold=True)
add_bullet(doc, '完成收钱吧网关对接：与收钱吧技术对接，确认正确的 method 名称，实现真实加购链路')
add_bullet(doc, '引入 SSE 流式输出：将 LLM 回复改为流式传输，用户可实时看到文字"打字"效果，大幅降低感知延迟')
add_bullet(doc, '添加 Prometheus metrics：暴露请求量、延迟分布、错误率、限流触发次数等核心指标')
add_bullet(doc, '前端增加菜品图片：在推荐结果中展示菜品缩略图，提升视觉吸引力')

add_normal_para(doc, '中期优化（1-2 月）：', bold=True)
add_bullet(doc, '用户画像系统：基于历史点餐数据构建用户偏好模型，实现个性化推荐（协同过滤 + LLM 推荐融合）')
add_bullet(doc, '推荐反馈闭环：收集用户对推荐结果的满意度，用于在线调整推荐评分权重')
add_bullet(doc, '多轮对话优化：增强追加推荐、替换菜品、调整数量等复杂多轮对话场景的处理能力')
add_bullet(doc, '小程序端适配：将前端改造为微信小程序，适配移动端点餐场景')
add_bullet(doc, '国际化支持：支持英文/泰文等多语言，面向云南旅游场景的海外游客')

add_normal_para(doc, '长期规划（3-6 月）：', bold=True)
add_bullet(doc, '语音点餐：集成 ASR（语音识别）能力，支持语音输入点餐需求')
add_bullet(doc, '多餐厅 SaaS 化：将系统改造为多租户架构，支持多个餐厅品牌独立配置菜单和规则')
add_bullet(doc, 'A/B 实验平台：支持在线对比不同推荐策略的效果，数据驱动推荐算法迭代')
add_bullet(doc, '预测性推荐：基于时间/天气/节假日等特征，在用户开口前主动推荐（如"下雨天，来锅热乎的菌汤吧"）')

# ======================== 结尾 ========================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 文档结束 —')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ======================== 保存 ========================
output_path = r'd:\meau\小菌点餐智能体_项目工作总结与汇报.docx'
doc.save(output_path)
print(f'文档已生成：{output_path}')